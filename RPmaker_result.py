from docx import Document
from docx.shared import Inches

# Create a new Word document
document = Document()

# Add a heading
document.add_heading('KR R2M 230223 업데이트 QA', level=1)

# Add test information as a table
table = document.add_table(rows=5, cols=2)
table.cell(0, 0).text = '프로젝트명'
table.cell(0, 1).text = 'KR R2M 230223 업데이트 QA'
table.cell(1, 0).text = '테스트 기간'
table.cell(1, 1).text = '2023.02.14 (화) ~ 2023.02.22 (수) (7D)'
table.cell(2, 0).text = '인원'
table.cell(2, 1).text = '박남주, 노진석, 성민석, 이재헌, 김영현, 황정수 (6명)'
table.cell(3, 0).text = '서버'
table.cell(3, 1).text = '[알파] 알파서버 그룹 01~04\n[라이브] i-Redcore 01'
table.cell(4, 0).text = '빌드 정보'
table.cell(4, 1).text = '(최종)R2MClientKorea_Alpha_150107_230216_2009.apk\n(최종)R2MClientKorea_Alpha_150107_230216_2009.ipa\n(최종)R2MClientKorea_Live_QA_150107_230217_1410.apk\n(최종)signed_gxs_R2MClientKorea_Live_150107_230217_1236_1676615021556.apk\n(최종)TestFlight 1.5.0 (107)'

# Save the document
document.save('KR_R2M_230223_update_QA.docx')